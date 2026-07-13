#!/usr/bin/env python3
"""
Compliance PRD Generator

Generates a complete Compliance PRD docx file from a compliance snapshot.
Implements team routing logic and applies exact formatting from the v2 template.
"""

import json
import sys
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===== CONFIGURATION =====
COLORS = {
    'dark_blue': RGBColor(31, 78, 121),      # 1F4E79 - headers
    'gray': RGBColor(90, 102, 117),          # 5A6675 - placeholder text
    'light_blue_bg': 'EAF1F8',               # light blue cell background
    'border_gray': 'C7D2DE',                 # border color
    'white': RGBColor(255, 255, 255),
}

FONT_NAME = 'Arial'
MARGIN_IN = 1.0  # 1 inch margins (1440 DXA)

# Team assignment heuristics
TEAM_KEYWORDS = {
    'Driver Ops': [
        'driver partner', 'delivery partner', 'independent contractor', 'driver',
        'gig worker', 'partner', 'acceptance', 'status', 'authentication',
        'background check', 'verification', 'onboarding'
    ],
    'Engineering': [
        'application', 'software', 'code', 'api', 'system', 'platform',
        'data', 'integration', 'database', 'algorithm', 'feature', 'deploy',
        'technical', 'backend', 'frontend', 'mobile app', 'driver app'
    ],
    'Legal': [
        'compliance', 'regulation', 'statute', 'law', 'requirement', 'legal',
        'binding', 'agreement', 'contract', 'disclosure', 'notice', 'audit'
    ]
}


def assign_team(requirement_text):
    """Assign requirement to appropriate team based on keyword matching."""
    text_lower = requirement_text.lower()

    # Count matches for each team
    scores = {}
    for team, keywords in TEAM_KEYWORDS.items():
        scores[team] = sum(1 for kw in keywords if kw in text_lower)

    # Return team with highest score, default to Legal
    best_team = max(scores, key=scores.get) if max(scores.values()) > 0 else 'Legal'
    return best_team


def set_cell_background(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_table_borders(table, color='000000', size='4'):
    """Add borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)

    tblPr.append(tblBorders)


def style_header_cell(cell, text):
    """Style a table header cell (dark blue background, white text)."""
    set_cell_background(cell, '1F4E79')
    paragraph = cell.paragraphs[0]
    paragraph.text = text
    for run in paragraph.runs:
        run.font.bold = True
        run.font.color.rgb = COLORS['white']
        run.font.size = Pt(10)


def style_placeholder_cell(cell, text):
    """Style a table data cell with placeholder text (gray italic)."""
    paragraph = cell.paragraphs[0]
    paragraph.text = text
    for run in paragraph.runs:
        run.font.italic = True
        run.font.color.rgb = COLORS['gray']
        run.font.size = Pt(10)


def add_heading_1(doc, text):
    """Add a Heading 1 with underline."""
    heading = doc.add_paragraph(text, style='Heading 1')
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(12)

    # Style the heading
    for run in heading.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLORS['dark_blue']

    # Add bottom border/underline
    pPr = heading._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return heading


def add_heading_2(doc, text):
    """Add a Heading 2."""
    heading = doc.add_paragraph(text, style='Heading 2')
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)

    for run in heading.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLORS['dark_blue']

    return heading


def add_body_text(doc, text):
    """Add body text paragraph."""
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)

    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(11)

    return para


def generate_prd(snapshot):
    """
    Generate a Compliance PRD docx from a compliance snapshot.

    Args:
        snapshot (dict): Compliance snapshot with keys:
            - matter_bill: str - bill/regulation short name
            - description: str - one-line description
            - jurisdiction: str - state/federal/EU
            - effective_date: str - YYYY-MM-DD
            - legal_owner: str - name
            - product_owner: str - name
            - statutory_trigger: str - law/rule citation
            - statutory_effective: str - when enforcement begins
            - in_scope: list - items in scope
            - out_of_scope: list - items out of scope
            - requirements: list - dicts with 'citation', 'current_state', 'future_state'
            - assumptions: list - assumptions
            - open_questions: list - open questions
            - references: list - reference links

    Returns:
        Document object
    """
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(MARGIN_IN)
        section.bottom_margin = Inches(MARGIN_IN)
        section.left_margin = Inches(MARGIN_IN)
        section.right_margin = Inches(MARGIN_IN)

    # ===== TITLE =====
    title = doc.add_paragraph('Compliance PRD')
    title.paragraph_format.space_after = Pt(6)
    for run in title.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = COLORS['dark_blue']

    # Subtitle
    subtitle = doc.add_paragraph(
        f"{snapshot.get('matter_bill', '[Bill / regulation short name]')} — {snapshot.get('description', '[one-line description of the product/ops change]')}"
    )
    subtitle.paragraph_format.space_after = Pt(12)
    for run in subtitle.runs:
        run.font.italic = True
        run.font.color.rgb = COLORS['gray']
        run.font.size = Pt(10)

    # ===== METADATA TABLE =====
    metadata_table = doc.add_table(rows=4, cols=4)
    metadata_table.autofit = False
    metadata_table.allow_autofit = False
    add_table_borders(metadata_table)

    # Set column widths (roughly equal)
    for row in metadata_table.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(1.3)
        row.cells[3].width = Inches(2.15)

    # Row 1: Matter/Bill, Jurisdiction
    style_header_cell(metadata_table.rows[0].cells[0], 'Matter / Bill')
    style_placeholder_cell(metadata_table.rows[0].cells[1], snapshot.get('matter_bill', '[short title]'))
    style_header_cell(metadata_table.rows[0].cells[2], 'Jurisdiction')
    style_placeholder_cell(metadata_table.rows[0].cells[3], snapshot.get('jurisdiction', '[state / federal / EU]'))
    set_cell_background(metadata_table.rows[0].cells[0], 'EAF1F8')
    set_cell_background(metadata_table.rows[0].cells[2], 'EAF1F8')

    # Row 2: Effective date, Status
    style_header_cell(metadata_table.rows[1].cells[0], 'Effective date')
    style_placeholder_cell(metadata_table.rows[1].cells[1], snapshot.get('effective_date', '[YYYY-MM-DD]'))
    style_header_cell(metadata_table.rows[1].cells[2], 'Status')
    style_placeholder_cell(metadata_table.rows[1].cells[3], '[Draft / In review / Final]')
    set_cell_background(metadata_table.rows[1].cells[0], 'EAF1F8')
    set_cell_background(metadata_table.rows[1].cells[2], 'EAF1F8')

    # Row 3: Legal owner, Product owner
    style_header_cell(metadata_table.rows[2].cells[0], 'Legal owner')
    style_placeholder_cell(metadata_table.rows[2].cells[1], snapshot.get('legal_owner', '[name]'))
    style_header_cell(metadata_table.rows[2].cells[2], 'Product owner')
    style_placeholder_cell(metadata_table.rows[2].cells[3], snapshot.get('product_owner', '[name]'))
    set_cell_background(metadata_table.rows[2].cells[0], 'EAF1F8')
    set_cell_background(metadata_table.rows[2].cells[2], 'EAF1F8')

    # Row 4: Date drafted, Version
    style_header_cell(metadata_table.rows[3].cells[0], 'Date drafted')
    style_placeholder_cell(metadata_table.rows[3].cells[1], snapshot.get('date_drafted', datetime.now().strftime('%Y-%m-%d')))
    style_header_cell(metadata_table.rows[3].cells[2], 'Version')
    style_placeholder_cell(metadata_table.rows[3].cells[3], 'v0.1')
    set_cell_background(metadata_table.rows[3].cells[0], 'EAF1F8')
    set_cell_background(metadata_table.rows[3].cells[2], 'EAF1F8')

    doc.add_paragraph()  # Spacing

    # ===== RACI MATRIX =====
    raci_table = doc.add_table(rows=2, cols=4)
    add_table_borders(raci_table)

    # Header row
    for i, label in enumerate(['Responsible\n(does the work)', 'Accountable\n(owns / approves)', 'Consulted\n(gives input)', 'Informed\n(kept in the loop)']):
        cell = raci_table.rows[0].cells[i]
        set_cell_background(cell, '1F4E79')
        para = cell.paragraphs[0]
        para.text = label
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = COLORS['white']
            run.font.size = Pt(10)

    # Data row
    raci_data = ['Product / Eng / Design lead', 'Legal owner', 'Privacy, Compliance, Ops', 'Exec sponsor, GM, support']
    for i, text in enumerate(raci_data):
        cell = raci_table.rows[1].cells[i]
        para = cell.paragraphs[0]
        para.text = text
        for run in para.runs:
            run.font.italic = True
            run.font.color.rgb = COLORS['gray']
            run.font.size = Pt(10)

    doc.add_paragraph()  # Spacing

    # ===== SECTION 1: SUMMARY =====
    add_heading_1(doc, '1. Summary: ')
    summary_text = snapshot.get('summary', '[2–3 sentences — what the law actually requires us to do]')
    add_body_text(doc, summary_text)

    # ===== SECTION 2: STATUTORY TRIGGER =====
    add_heading_1(doc, '2. Statutory trigger')

    law_para = doc.add_paragraph()
    law_run = law_para.add_run('Law / rule: ')
    law_run.bold = True
    law_run.font.name = FONT_NAME
    law_text_run = law_para.add_run(snapshot.get('statutory_trigger', '[e.g., California AB-XXXX, FMCSA § 395, NYC Local Law 144]'))
    law_text_run.italic = True
    law_text_run.font.color.rgb = COLORS['gray']
    law_text_run.font.size = Pt(10)

    eff_para = doc.add_paragraph()
    eff_run = eff_para.add_run('Effective date: ')
    eff_run.bold = True
    eff_run.font.name = FONT_NAME
    eff_text_run = eff_para.add_run(f"enforcement begins {snapshot.get('statutory_effective', '[YYYY-MM-DD]')}")
    eff_text_run.italic = True
    eff_text_run.font.color.rgb = COLORS['gray']
    eff_text_run.font.size = Pt(10)

    # ===== SECTION 3: SCOPE =====
    add_heading_1(doc, '3. Scope')

    add_heading_2(doc, 'In scope')
    for item in snapshot.get('in_scope', ['[surfaces, products, regions, user cohorts covered]']):
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_2(doc, 'Out of scope')
    for item in snapshot.get('out_of_scope', ['[explicitly excluded surfaces — e.g., UK/EU, BevMo, mobile app phase 2]']):
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    # ===== SECTION 4: REQUIREMENTS =====
    add_heading_1(doc, '4. Statutory requirements → current state → future state')

    intro = doc.add_paragraph('One row per discrete statutory obligation. Cite the section. Be specific about gap and future-state requirement so engineering can act on it.')
    intro.paragraph_format.space_after = Pt(6)
    for run in intro.runs:
        run.font.italic = True
        run.font.color.rgb = COLORS['gray']
        run.font.size = Pt(10)

    # Requirements table
    req_table = doc.add_table(rows=1, cols=4)
    add_table_borders(req_table)

    # Header row
    headers = ['Statutory requirement\n(cite)', 'Current state', 'Gap / risk', 'Future-state requirement']
    for i, header in enumerate(headers):
        cell = req_table.rows[0].cells[i]
        set_cell_background(cell, '1F4E79')
        para = cell.paragraphs[0]
        para.text = header
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = COLORS['white']
            run.font.size = Pt(10)

    # Add requirement rows
    requirements = snapshot.get('requirements', [])
    if not requirements:
        # Add example row
        requirements = [{
            'citation': 'e.g., § 1798.135(b)(1) — honor GPC signal',
            'current_state': '[What we do today (1–2 sentences)]',
            'gap_risk': '[Compliant / Partial / Non-compliant + risk]',
            'future_state': '[Team name]: specific action/requirement'
        }]

    for req in requirements:
        row = req_table.add_row()

        # Statutory requirement
        citation_cell = row.cells[0]
        citation_para = citation_cell.paragraphs[0]
        citation_run = citation_para.add_run('Req ')
        citation_run.bold = True
        citation_para.add_run('\n')
        citation_text_run = citation_para.add_run(req.get('citation', ''))
        citation_text_run.italic = True
        citation_text_run.font.color.rgb = COLORS['gray']
        citation_text_run.font.size = Pt(10)

        # Current state
        current_cell = row.cells[1]
        current_para = current_cell.paragraphs[0]
        current_run = current_para.add_run(req.get('current_state', ''))
        current_run.italic = True
        current_run.font.color.rgb = COLORS['gray']
        current_run.font.size = Pt(10)

        # Gap / risk
        gap_cell = row.cells[2]
        gap_para = gap_cell.paragraphs[0]
        gap_run = gap_para.add_run(req.get('gap_risk', ''))
        gap_run.italic = True
        gap_run.font.color.rgb = COLORS['gray']
        gap_run.font.size = Pt(10)

        # Future-state requirement (with team assignment)
        future_cell = row.cells[3]
        future_para = future_cell.paragraphs[0]
        future_text = req.get('future_state', '')

        # Assign team if not already in text
        if not any(team in future_text for team in ['Driver Ops', 'Engineering', 'Legal']):
            assigned_team = assign_team(req.get('citation', '') + ' ' + req.get('future_state', ''))
            future_text = f"{assigned_team}: {future_text}"

        # Parse team name and action
        if ':' in future_text:
            team_part, action_part = future_text.split(':', 1)
            team_run = future_para.add_run(team_part.strip())
            team_run.bold = True
            team_run.underline = True
            team_run.font.color.rgb = COLORS['gray']
            team_run.font.size = Pt(10)

            action_run = future_para.add_run(f":{action_part}")
            action_run.italic = True
            action_run.font.color.rgb = COLORS['gray']
            action_run.font.size = Pt(10)
        else:
            future_run = future_para.add_run(future_text)
            future_run.italic = True
            future_run.font.color.rgb = COLORS['gray']
            future_run.font.size = Pt(10)

    # ===== SECTION 5: ASSUMPTIONS & QUESTIONS =====
    add_heading_1(doc, '5. Assumptions & open questions')

    add_heading_2(doc, 'Assumptions')
    for item in snapshot.get('assumptions', ['[...]']):
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    add_heading_2(doc, 'Open questions for legal / regulator')
    for item in snapshot.get('open_questions', ['[...]']):
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    # ===== SECTION 6: IMPLEMENTATION PLAN =====
    add_heading_1(doc, '6. Implementation plan')

    intro = doc.add_paragraph('P0 = required by effective date. P1 = required but flexible on timing. P2 = nice-to-have / hardening.')
    intro.paragraph_format.space_after = Pt(6)
    for run in intro.runs:
        run.font.italic = True
        run.font.color.rgb = COLORS['gray']
        run.font.size = Pt(10)

    impl_table = doc.add_table(rows=1, cols=6)
    add_table_borders(impl_table)

    impl_headers = ['Pri', '#', 'Description', 'Owner / team', 'Sizing', 'Target']
    for i, header in enumerate(impl_headers):
        cell = impl_table.rows[0].cells[i]
        set_cell_background(cell, '1F4E79')
        para = cell.paragraphs[0]
        para.text = header
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = COLORS['white']
            run.font.size = Pt(10)

    # Add placeholder rows
    for pri, num in [('P0', '1'), ('P0', '2'), ('P1', '3'), ('P1', '4'), ('P2', '5')]:
        row = impl_table.add_row()
        for i, text in enumerate([pri, num, '[work item]', '[name]', '[S/M/L]', '[YYYY-MM-DD]']):
            cell = row.cells[i]
            para = cell.paragraphs[0]
            run = para.add_run(text)
            if '[' in text:
                run.italic = True
                run.font.color.rgb = COLORS['gray']
            run.font.size = Pt(10)

    # ===== SECTION 7: SUCCESS CRITERIA =====
    add_heading_1(doc, '7. Success criteria & compliance verification')

    criteria_items = [
        ('Acceptance test:', '[...]'),
        ('Metric:', '[...]'),
        ('Sign-off:', '[...]')
    ]
    for label, placeholder in criteria_items:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{label} ')
        run1.bold = True
        run1.font.name = FONT_NAME
        run2 = p.add_run(placeholder)
        run2.italic = True
        run2.font.color.rgb = COLORS['gray']
        run2.font.size = Pt(10)

    # ===== SECTION 8: REFERENCES =====
    add_heading_1(doc, '8. References')

    for ref in snapshot.get('references', ['[link to statute / regulation text]', '[link to enforcement guidance, FAQ, or AG advisory]', '[link to internal Jira epic, design doc, or related PRD]']):
        p = doc.add_paragraph(ref, style='List Bullet')
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(11)

    return doc


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python generate_compliance_prd.py <snapshot.json> [output.docx]")
        sys.exit(1)

    snapshot_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else 'Compliance_PRD.docx'

    # Load snapshot
    with open(snapshot_file) as f:
        snapshot = json.load(f)

    # Generate document
    doc = generate_prd(snapshot)
    doc.save(output_file)

    print(f"✓ Generated {output_file}")
